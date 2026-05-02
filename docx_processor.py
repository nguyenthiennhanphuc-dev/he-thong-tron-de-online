import io
import zipfile
import random
import re
import copy
from docx import Document
from docxcompose.composer import Composer
from docx.shared import RGBColor

def fix_docx_floats(file_stream):
    """Gọt số thập phân chống crash"""
    file_stream.seek(0)
    zip_in = zipfile.ZipFile(file_stream)
    zip_buffer = io.BytesIO()
    zip_out = zipfile.ZipFile(zip_buffer, 'w')
    for item in zip_in.infolist():
        content = zip_in.read(item.filename)
        if item.filename.endswith('.xml'):
            content_str = content.decode('utf-8')
            header_end = content_str.find('?>') + 2
            if header_end > 1:
                header = content_str[:header_end]
                body = content_str[header_end:]
                body = re.sub(r'="(\-?[0-9]+)\.[0-9]+"', r'="\1"', body)
                content_str = header + body
            else:
                content_str = re.sub(r'="(\-?[0-9]+)\.[0-9]+"', r'="\1"', content_str)
            content = content_str.encode('utf-8')
        zip_out.writestr(item, content)
    zip_out.close()
    zip_buffer.seek(0)
    return zip_buffer

def merge_and_track_files(files):
    """Hợp thể các file để không làm đứt liên kết kho ảnh"""
    master_stream = fix_docx_floats(files[0])
    master = Document(master_stream)
    p_src = master.add_paragraph(f"@SOURCE_FILE: {files[0].name}")
    master._body._body.insert(0, p_src._element) 
    composer = Composer(master)
    for f in files[1:]:
        f_stream = fix_docx_floats(f)
        doc_append = Document(f_stream)
        p_src_app = doc_append.add_paragraph(f"@SOURCE_FILE: {f.name}")
        doc_append._body._body.insert(0, p_src_app._element)
        composer.append(doc_append)
    merged_io = io.BytesIO()
    composer.doc.save(merged_io)
    merged_io.seek(0)
    return merged_io

def clean_source_tags(paragraph):
    """Tự động dọn dẹp các mã nguồn như [1], """
    for run in paragraph.runs:
        if run.text: run.text = re.sub(r'\[.*?\]\s*', '', run.text)

def enhance_visuals(paragraph):
    """Phóng to chữ cho học sinh dễ đọc và căn giữa hình ảnh"""
    for run in paragraph.runs:
        if run.font.size is None: run.font.size = 177800
    if 'Graphic' in paragraph._p.xml or 'pic:pic' in paragraph._p.xml:
        paragraph.alignment = 1

def extract_sections_xml(doc):
    """Máy quét phân khu (Phần I, II, III), Nhóm chung và Đọc chữ HẾT"""
    sections = []
    current_section = {"title": "CHUNG", "intro": [], "items": []} 
    current_q = []
    current_g = None
    outro = []
    is_outro = False
    current_src = "Chưa rõ nguồn"
    is_template = True
    file_count = 0
    
    for p in doc.paragraphs:
        clean_source_tags(p)
        enhance_visuals(p)
        text = p.text.strip()
        
        if text.startswith("@SOURCE_FILE:"):
            if current_q:
                if current_g: current_g["questions"].append(current_q)
                else: current_section["items"].append({"type": "q", "data": current_q, "src": current_src, "is_temp": is_template})
                current_q = []
            if current_g:
                current_section["items"].append({"type": "g", "data": current_g, "src": current_src, "is_temp": is_template})
                current_g = None
            if current_section["intro"] or current_section["items"]:
                sections.append(current_section)
            current_section = {"title": "CHUNG", "intro": [], "items": []}
            is_outro = False
            current_src = text.replace("@SOURCE_FILE:", "").strip()
            file_count += 1
            if file_count > 1: is_template = False
            p.text = "" 
            continue
            
        if is_outro:
            if is_template: outro.append(p)
            continue
            
        is_het = re.match(r'^[\-\—\_\s\*]*HẾT[\-\—\_\s\*]*$', text.upper())
        is_phan = re.match(r'^PHẦN\s+[IVX\d]+', text.upper())
        is_start_group = "@BẮT ĐẦU" in text.upper()
        is_end_group = "@KẾT THÚC" in text.upper()
        is_cau = re.match(r'^Câu\s+\d+[\.:]', text, re.IGNORECASE)
        
        if is_het:
            is_outro = True
            if current_q:
                if current_g: current_g["questions"].append(current_q)
                else: current_section["items"].append({"type": "q", "data": current_q, "src": current_src, "is_temp": is_template})
                current_q = []
            if current_g:
                current_section["items"].append({"type": "g", "data": current_g, "src": current_src, "is_temp": is_template})
                current_g = None
            if current_section["intro"] or current_section["items"]:
                sections.append(current_section)
                current_section = {"title": "CHUNG", "intro": [], "items": []}
            if is_template: outro.append(p)
            continue
        
        if is_phan:
            if current_q:
                if current_g: current_g["questions"].append(current_q)
                else: current_section["items"].append({"type": "q", "data": current_q, "src": current_src, "is_temp": is_template})
                current_q = []
            if current_g:
                current_section["items"].append({"type": "g", "data": current_g, "src": current_src, "is_temp": is_template})
                current_g = None
            if current_section["intro"] or current_section["items"]:
                sections.append(current_section)
            current_section = {"title": text, "intro": [p], "items": []}
        elif is_start_group:
            if current_q:
                if current_g: current_g["questions"].append(current_q)
                else: current_section["items"].append({"type": "q", "data": current_q, "src": current_src, "is_temp": is_template})
                current_q = []
            if current_g: current_section["items"].append({"type": "g", "data": current_g, "src": current_src, "is_temp": is_template})
            current_g = {"intro": [p], "questions": []}
        elif is_end_group:
            if current_g:
                current_g["intro"].append(p)
                if current_q:
                    current_g["questions"].append(current_q)
                    current_q = []
                current_section["items"].append({"type": "g", "data": current_g, "src": current_src, "is_temp": is_template})
                current_g = None
            else:
                if current_q: current_q.append(p)
                else: current_section["intro"].append(p)
        elif is_cau:
            if current_q:
                if current_g: current_g["questions"].append(current_q)
                else: current_section["items"].append({"type": "q", "data": current_q, "src": current_src, "is_temp": is_template})
            current_q = [p]
        else:
            if current_q: current_q.append(p)
            elif current_g: current_g["intro"].append(p)
            else: current_section["intro"].append(p)
            
    if current_q:
        if current_g: current_g["questions"].append(current_q)
        else: current_section["items"].append({"type": "q", "data": current_q, "src": current_src, "is_temp": is_template})
    if current_g:
        current_section["items"].append({"type": "g", "data": current_g, "src": current_src, "is_temp": is_template})
    if current_section["intro"] or current_section["items"]:
        sections.append(current_section)
    return sections, outro

def fix_choice_label(paragraph, new_label):
    try:
        for run in paragraph.runs:
            text = run.text.strip()
            if re.match(r'^[A-D]\*?[\.\)]\*?', text):
                run.text = re.sub(r'^[A-D]\*?[\.\)]\*?', new_label, run.text, count=1)
                break
            elif text in ['A', 'B', 'C', 'D', 'A*', 'B*', 'C*', 'D*']:
                run.text = re.sub(r'^[A-D]\*?', new_label.replace('.', ''), run.text, count=1)
                break
    except Exception: pass

def clean_asterisk_label(paragraph):
    try:
        for run in paragraph.runs:
            text = run.text.lstrip()
            if re.match(r'^[a-d]\*?[\.\)]\*?', text):
                run.text = re.sub(r'^(\s*[a-d])\*?([\.\)])\*?', r'\1\2', run.text, count=1)
                break
            elif text in ['a', 'b', 'c', 'd', 'a*', 'b*', 'c*', 'd*']:
                run.text = run.text.replace('*', '')
                break
    except Exception: pass

def is_correct_choice(paragraph):
    """
    Sniper Model v4: Quét toàn bộ Run bằng XML thuần túy để đảm bảo bắt dính 100% định dạng.
    """
    try:
        # Kiểm tra thuộc tính paragraph cho auto-numbering
        p_xml = paragraph._p.xml
        if '<w:pPr>' in p_xml:
            pPr_xml = p_xml.split('<w:pPr>')[1].split('</w:pPr>')[0]
            if 'w:val="FF0000"' in pPr_xml or 'w:val="red"' in pPr_xml.lower() or 'w:val="C00000"' in pPr_xml or '<w:u ' in pPr_xml or '<w:u/>' in pPr_xml or '<w:highlight' in pPr_xml:
                return True
    except Exception: pass

    try:
        for run in paragraph.runs:
            if run.text.strip(): 
                xml_str = run._element.xml
                if 'w:val="FF0000"' in xml_str or 'w:val="red"' in xml_str.lower() or 'w:val="C00000"' in xml_str or '<w:u ' in xml_str or '<w:u/>' in xml_str or '<w:highlight' in xml_str:
                    return True
                if run.font.color and (str(run.font.color.rgb) == 'FF0000' or str(run.font.color.rgb) == 'C00000'):
                    return True
                if run.font.underline:
                    return True
    except Exception: pass
    
    return False

def extract_answer_text(paragraph):
    """Hút phần nội dung chữ màu đỏ/gạch chân (Dự phòng)"""
    ans_parts = []
    for run in paragraph.runs:
        try:
            xml_str = run._element.xml
            if 'w:val="FF0000"' in xml_str or 'w:val="red"' in xml_str.lower() or 'w:val="C00000"' in xml_str or '<w:u ' in xml_str or '<w:u/>' in xml_str:
                if run.text: ans_parts.append(run.text)
        except Exception: pass
    return "".join(ans_parts).strip()

def process_question_block(q_blocks, q_num, shuffle_a, keep_color=False, manual_ans=None):
    """Trái tim xử lý: Nhận diện, Đổi màu tàng hình, Rút trích Đ/S và Lời giải"""
    q_blocks = copy.deepcopy(q_blocks)
    first_p = q_blocks[0]
    first_p.text = re.sub(r'^Câu\s+\d+', f'Câu {q_num}', first_p.text, flags=re.IGNORECASE)
    
    correct_label = ""
    mcq_indices = []
    tf_indices = []
    numbered_indices = []
    
    for i, p in enumerate(q_blocks):
        text = p.text.strip()
        if re.match(r'^[A-D]\*?[\.\)]\*?', text): mcq_indices.append(i)
        elif re.match(r'^[a-d]\*?[\.\)]\*?', text): tf_indices.append(i)
        elif '<w:numPr>' in p._p.xml: numbered_indices.append(i)
        
    elements_to_keep = []
    
    if len(mcq_indices) == 0 and len(numbered_indices) == 4:
        mcq_indices = numbered_indices
        
    if len(mcq_indices) == 4:
        choices = [{"p": q_blocks[i], "orig_label": l} for i, l in zip(mcq_indices, ['A', 'B', 'C', 'D'])]
        if shuffle_a: random.shuffle(choices)
        labels = ['A', 'B', 'C', 'D']; labels_text = ['A.', 'B.', 'C.', 'D.']
        for i, orig_idx in enumerate(mcq_indices):
            choice_dict = choices[i]
            p = choice_dict["p"]
            orig_l = choice_dict["orig_label"]
            
            has_asterisk = bool(re.match(r'^[A-D]\*?[\.\)]\*?', p.text.strip()) and '*' in p.text.strip().split(' ', 1)[0])
            
            is_correct = False
            if manual_ans and isinstance(manual_ans, str) and manual_ans.upper() == orig_l:
                is_correct = True
            elif not manual_ans and (is_correct_choice(p) or has_asterisk):
                is_correct = True
                
            new_label_text = labels_text[i]
            if is_correct: 
                correct_label = labels[i]
                if keep_color: new_label_text = new_label_text.replace('.', '*.').replace(')', ')*')
                
            q_blocks[orig_idx] = p
            fix_choice_label(q_blocks[orig_idx], new_label_text)
        if not correct_label and not shuffle_a:
            c_idx = 0
            for p in q_blocks:
                if re.match(r'^[A-D][\.\)]', p.text.strip()):
                    if is_correct_choice(p) and c_idx < 4: correct_label = labels[c_idx]
                    c_idx += 1
                    
        for p in q_blocks:
            if not keep_color:
                for run in p.runs:
                    try:
                        xml_str = run._element.xml
                        if 'w:val="FF0000"' in xml_str or 'w:val="red"' in xml_str.lower() or 'w:val="C00000"' in xml_str or '<w:u ' in xml_str or '<w:u/>' in xml_str or '<w:highlight' in xml_str:
                            run.font.color.rgb = RGBColor(0, 0, 0)
                            run.font.underline = False
                            run.font.bold = False
                    except Exception: pass
            elements_to_keep.append(copy.deepcopy(p._element))
            
    elif len(tf_indices) == 4:
        tf_answers = []
        for i, tf_idx in enumerate(tf_indices):
            ans = "S"
            p = q_blocks[tf_idx]
            has_asterisk = bool(re.match(r'^[a-d]\*?[\.\)]\*?', p.text.strip()) and '*' in p.text.strip().split(' ', 1)[0])
            
            if manual_ans and isinstance(manual_ans, list) and len(manual_ans) == 4:
                ans = manual_ans[i]
            elif not manual_ans and (is_correct_choice(p) or has_asterisk): 
                ans = "Đ"
            tf_answers.append(ans)
            
            if not keep_color:
                clean_asterisk_label(p)
            
        correct_label = ", ".join([f"{['a','b','c','d'][idx]}-{ans}" for idx, ans in enumerate(tf_answers)])
        
        for p in q_blocks:
            if not keep_color:
                for run in p.runs:
                    try:
                        xml_str = run._element.xml
                        if 'w:val="FF0000"' in xml_str or 'w:val="red"' in xml_str.lower() or '<w:u ' in xml_str or '<w:u/>' in xml_str:
                            run.font.color.rgb = RGBColor(0, 0, 0)
                            run.font.underline = False
                            run.font.bold = False
                    except Exception: pass
            elements_to_keep.append(copy.deepcopy(p._element))
            
    else:
        ans_texts = []
        if manual_ans and isinstance(manual_ans, str):
            correct_label = manual_ans
            for p in q_blocks:
                if not keep_color:
                    for run in p.runs:
                        try:
                            xml_str = run._element.xml
                            if 'w:val="FF0000"' in xml_str or 'w:val="red"' in xml_str.lower() or 'w:val="C00000"' in xml_str or '<w:u ' in xml_str or '<w:u/>' in xml_str or '<w:highlight' in xml_str:
                                run.font.color.rgb = RGBColor(0, 0, 0)
                                run.font.underline = False
                        except Exception: pass
                elements_to_keep.append(copy.deepcopy(p._element))
        else:
            for p in q_blocks:
                text_strip = p.text.strip()
                
                match_prefix = re.match(r'^(Đáp án|ĐS|Kết quả|Lời giải|Đáp số)[\s\:\.]*(.*)', text_strip, re.IGNORECASE)
                if match_prefix:
                    ans_val = match_prefix.group(2).strip()
                    if ans_val: ans_texts.append(ans_val)
                    if not keep_color:
                        continue 
                    
                txt_red = extract_answer_text(p)
                if txt_red:
                    clean_txt = re.sub(r'^(Đáp án|ĐS|Kết quả|Lời giải|Đáp số)[\s\:]*', '', txt_red, flags=re.IGNORECASE).strip()
                    if clean_txt: ans_texts.append(clean_txt)
                    if not keep_color:
                        for run in p.runs:
                            try:
                                xml_str = run._element.xml
                                if 'w:val="FF0000"' in xml_str or 'w:val="red"' in xml_str.lower() or '<w:u ' in xml_str or '<w:u/>' in xml_str:
                                    run.text = "" 
                            except Exception: pass
                        
                elements_to_keep.append(copy.deepcopy(p._element))
                
            if ans_texts: correct_label = " ; ".join(ans_texts)
            
    return correct_label, elements_to_keep

def allocate_questions_fairly(pool, needed, randomize=True):
    """
    Thuật toán chia đều số lượng câu hỏi từ nhiều nguồn (files) khác nhau.
    Có khả năng bù trừ (nếu file A thiếu câu thì bốc bù từ file B).
    randomize=True: bốc ngẫu nhiên. False: lấy theo thứ tự gốc trong file.
    """
    pool_by_src = {}
    for item in pool:
        src = item["src"]
        if src not in pool_by_src:
            pool_by_src[src] = []
        pool_by_src[src].append(item)
        
    sources = list(pool_by_src.keys())
    if not sources or needed <= 0:
        return []
        
    quota = {src: 0 for src in sources}
    base = needed // len(sources)
    rem = needed % len(sources)
    
    # Bước 1: Chia đều cơ bản
    for src in sources:
        quota[src] = base
    # Phân phát số dư (nếu chia không hết)
    if randomize:
        for src in random.sample(sources, rem):
            quota[src] += 1
    else:
        for i in range(rem):
            quota[sources[i]] += 1
        
    # Bước 2: Bù trừ (Nếu có file không đủ số lượng quota yêu cầu)
    while True:
        shortfall = 0
        active_sources = []
        for src in sources:
            avail = len(pool_by_src[src])
            if quota[src] > avail:
                shortfall += (quota[src] - avail)
                quota[src] = avail # Chốt lấy tối đa số câu file này có
            elif quota[src] < avail:
                active_sources.append(src) # Các file còn dư dả để bốc bù
        
        if shortfall == 0 or not active_sources:
            break # Hoàn hảo hoặc hết file để bù
            
        add_base = shortfall // len(active_sources)
        add_rem = shortfall % len(active_sources)
        for src in active_sources:
            quota[src] += add_base
        if randomize:
            for src in random.sample(active_sources, add_rem):
                quota[src] += 1
        else:
            for i in range(add_rem):
                quota[active_sources[i]] += 1

    # Bước 3: Bốc câu hỏi dựa trên Quota đã chốt
    selected_items = []
    for src in sources:
        if quota[src] > 0:
            count = min(quota[src], len(pool_by_src[src]))
            if randomize:
                selected_items.extend(random.sample(pool_by_src[src], count))
            else:
                # Lấy theo thứ tự gốc trong file (không xáo trộn)
                selected_items.extend(pool_by_src[src][:count])
            
    return selected_items



def normalize_clumped_choices(doc):
    import re
    from copy import deepcopy
    pattern = re.compile(r'(\s+)([A-Da-d]\*?[\.\)])(\s+)')
    
    for p in tuple(doc.paragraphs):
        text = p.text
        if not text.strip(): continue
        
        matches = list(pattern.finditer(text))
        if not matches: continue
        
        split_indices = [0] + [m.start(2) for m in matches] + [len(text)]
        
        for i in range(len(split_indices) - 1):
            start_idx = split_indices[i]
            end_idx = split_indices[i+1]
            
            if start_idx == end_idx: continue
            
            new_p = p.insert_paragraph_before()
            if p._p.pPr is not None: new_p._p.append(deepcopy(p._p.pPr))
            
            current_len = 0
            for run in p.runs:
                if not run.text:
                    if current_len >= start_idx and current_len < end_idx:
                        new_p._p.append(deepcopy(run._element))
                    continue
                    
                run_text = run.text
                run_len = len(run_text)
                
                overlap_start = max(current_len, start_idx)
                overlap_end = min(current_len + run_len, end_idx)
                
                if overlap_start < overlap_end:
                    rel_start = overlap_start - current_len
                    rel_end = overlap_end - current_len
                    segment_text = run_text[rel_start:rel_end]
                    
                    new_run = deepcopy(run._element)
                    for t in new_run.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                        t.text = segment_text
                    new_p._p.append(new_run)
                    
                current_len += run_len
                
        p._element.getparent().remove(p._element)

def parse_and_build_pools(files):
    merged_stream = merge_and_track_files(files)
    doc = Document(merged_stream)
    
    # Chuẩn hóa các phương án dính liền nhau (A. B. C. D. trên 1 dòng)
    normalize_clumped_choices(doc)
    
    sections, outro = extract_sections_xml(doc)
    
    global_pools = {} 
    structure = []
    
    title_to_part = {}
    part_idx = 0
    
    for sec in sections:
        title = sec['title']
        if title != 'CHUNG' and title not in title_to_part:
            match = re.search(r'PHẦN\s+([IVX]+)', title.upper())
            if match:
                roman = match.group(1)
                if roman == 'I': title_to_part[title] = 'part_1'
                elif roman == 'II': title_to_part[title] = 'part_2'
                elif roman == 'III': title_to_part[title] = 'part_3'
                else: 
                    part_idx += 1
                    title_to_part[title] = f'part_{part_idx}'
            else:
                part_idx += 1
                title_to_part[title] = f'part_{part_idx}'
            
        if title not in global_pools: global_pools[title] = []
        
        start_q_idx = sum(1 for item in global_pools[title] if item['type'] == 'q')
        start_q_idx += sum(len(item['data']['questions']) for item in global_pools[title] if item['type'] == 'g')
        
        current_q_idx = start_q_idx + 1
        
        for item in sec['items']:
            if item['type'] == 'q':
                item['original_index'] = current_q_idx
                current_q_idx += 1
            elif item['type'] == 'g':
                item['original_indices'] = list(range(current_q_idx, current_q_idx + len(item['data']['questions'])))
                current_q_idx += len(item['data']['questions'])
                
        global_pools[title].extend(sec['items'])
        if len(sec['items']) > 0 and sec['items'][0]['is_temp']:
            if not any(s['title'] == title for s in structure):
                structure.append({'title': title, 'intro': sec['intro'], 'needed': len(sec['items'])})
            else:
                for s in structure:
                    if s['title'] == title: s['needed'] += len(sec['items'])

    return global_pools, structure, title_to_part, outro, merged_stream.getvalue()

def generate_preview(global_pools, structure, title_to_part, manual_answers=None):
    preview_data = []
    
    q_count_mock = 1
    for sec_template in structure:
        title = sec_template['title']
        pool = global_pools.get(title, [])
        part_key = title_to_part.get(title)
        part_manual_ans = manual_answers.get(part_key, {}) if manual_answers else {}
        
        part_preview = {'title': title, 'questions': []}
        
        for item in pool:
            if item['type'] == 'q':
                manual_ans_for_item = part_manual_ans.get(item.get('original_index'))
                correct_label, els = process_question_block(item['data'], q_count_mock, shuffle_a=False, keep_color=True, manual_ans=manual_ans_for_item)
                
                full_text = '\n'.join(p.text for p in els if p.text.strip())
                part_preview['questions'].append({
                    'text': full_text, 
                    'answer': correct_label,
                    'orig_idx': item.get('original_index')
                })
                q_count_mock += 1
            elif item['type'] == 'g':
                g_data = copy.deepcopy(item['data'])
                intro_text = '\n'.join(p.text for p in g_data['intro'] if p.text.strip())
                
                group_qs = g_data['questions']
                sub_q_indices = item.get('original_indices', [])
                zipped_qs = list(zip(group_qs, sub_q_indices))
                
                for q_blocks, orig_sub_idx in zipped_qs:
                    manual_ans_for_sub = part_manual_ans.get(orig_sub_idx)
                    correct_label, els = process_question_block(q_blocks, q_count_mock, shuffle_a=False, keep_color=True, manual_ans=manual_ans_for_sub)
                    full_text = '\n'.join(p.text for p in els if p.text.strip())
                    
                    part_preview['questions'].append({
                        'text': f'[Nhóm dùng chung]\n{intro_text}\n{full_text}', 
                        'answer': correct_label,
                        'orig_idx': orig_sub_idx
                    })
                    q_count_mock += 1
                    
        preview_data.append(part_preview)
    return preview_data

def build_answer_tables(doc, ans_by_part, ma_de=None):
    import re
    if ma_de:
        doc.add_paragraph(f'BẢNG ĐÁP ÁN - MÃ ĐỀ {ma_de}').bold = True
        
    p1_list = ans_by_part.get('part_1', [])
    if p1_list:
        doc.add_paragraph('BẢNG ĐÁP ÁN PHẦN 1').bold = True
        chunk_size = 15
        for i in range(0, len(p1_list), chunk_size):
            chunk = p1_list[i:i+chunk_size]
            table = doc.add_table(rows=2, cols=len(chunk) + 1)
            table.style = 'Table Grid'
            table.cell(0,0).text = 'Câu'
            table.cell(1,0).text = 'Đáp án'
            for j, item in enumerate(chunk):
                table.cell(0, j+1).text = str(item['q'])
                table.cell(1, j+1).text = str(item['ans'])
        doc.add_paragraph('')

    p2_list = ans_by_part.get('part_2', [])
    if p2_list:
        doc.add_paragraph('BẢNG ĐÁP ÁN PHẦN 2').bold = True
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Câu'; hdr[1].text = 'a'; hdr[2].text = 'b'; hdr[3].text = 'c'; hdr[4].text = 'd'
        
        for item in p2_list:
            row = table.add_row().cells
            row[0].text = str(item['q'])
            ans_str = str(item['ans'])
            vals = {'a': '', 'b': '', 'c': '', 'd': ''}
            matches = re.findall(r'([a-d])\s*-\s*([ĐS])', ans_str, re.IGNORECASE)
            for k, v in matches: vals[k.lower()] = v.upper()
            row[1].text = vals['a']; row[2].text = vals['b']; row[3].text = vals['c']; row[4].text = vals['d']
        doc.add_paragraph('')

    p3_list = ans_by_part.get('part_3', [])
    if p3_list:
        doc.add_paragraph('BẢNG ĐÁP ÁN PHẦN 3').bold = True
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Câu'; hdr[1].text = 'Đáp án'
        for item in p3_list:
            row = table.add_row().cells
            row[0].text = str(item['q'])
            row[1].text = str(item['ans']).replace(';', '\n')
        doc.add_paragraph('')
        
    other_list = ans_by_part.get('other', [])
    for k, lst in ans_by_part.items():
        if k not in ['part_1', 'part_2', 'part_3', 'other']:
            other_list.extend(lst)
            
    if other_list:
        doc.add_paragraph('BẢNG ĐÁP ÁN KHÁC').bold = True
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Câu'; hdr[1].text = 'Đáp án'; hdr[2].text = 'Câu'; hdr[3].text = 'Đáp án'
        total_qs = len(other_list)
        for i in range(0, total_qs, 2):
            row = table.add_row().cells
            row[0].text = str(other_list[i]['q']); row[1].text = str(other_list[i]['ans'])
            if i + 1 < total_qs:
                row[2].text = str(other_list[i+1]['q']); row[3].text = str(other_list[i+1]['ans'])
        doc.add_paragraph('')

def build_answer_tables(doc, ans_by_part, ma_de=None):
    import re
    if ma_de:
        doc.add_paragraph(f'BẢNG ĐÁP ÁN - MÃ ĐỀ {ma_de}').bold = True
        
    p1_list = ans_by_part.get('part_1', [])
    if p1_list:
        doc.add_paragraph('BẢNG ĐÁP ÁN PHẦN 1').bold = True
        chunk_size = 15
        for i in range(0, len(p1_list), chunk_size):
            chunk = p1_list[i:i+chunk_size]
            table = doc.add_table(rows=2, cols=len(chunk) + 1)
            table.style = 'Table Grid'
            table.cell(0,0).text = 'Câu'
            table.cell(1,0).text = 'Đáp án'
            for j, item in enumerate(chunk):
                table.cell(0, j+1).text = str(item['q'])
                table.cell(1, j+1).text = str(item['ans'])
        doc.add_paragraph('')

    p2_list = ans_by_part.get('part_2', [])
    if p2_list:
        doc.add_paragraph('BẢNG ĐÁP ÁN PHẦN 2').bold = True
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Câu'; hdr[1].text = 'a'; hdr[2].text = 'b'; hdr[3].text = 'c'; hdr[4].text = 'd'
        
        for item in p2_list:
            row = table.add_row().cells
            row[0].text = str(item['q'])
            ans_str = str(item['ans'])
            vals = {'a': '', 'b': '', 'c': '', 'd': ''}
            matches = re.findall(r'([a-d])\s*-\s*([ĐS])', ans_str, re.IGNORECASE)
            for k, v in matches: vals[k.lower()] = v.upper()
            row[1].text = vals['a']; row[2].text = vals['b']; row[3].text = vals['c']; row[4].text = vals['d']
        doc.add_paragraph('')

    p3_list = ans_by_part.get('part_3', [])
    if p3_list:
        doc.add_paragraph('BẢNG ĐÁP ÁN PHẦN 3').bold = True
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Câu'; hdr[1].text = 'Đáp án'
        for item in p3_list:
            row = table.add_row().cells
            row[0].text = str(item['q'])
            row[1].text = str(item['ans']).replace(';', '\n')
        doc.add_paragraph('')
        
    other_list = ans_by_part.get('other', [])
    for k, lst in ans_by_part.items():
        if k not in ['part_1', 'part_2', 'part_3', 'other']:
            other_list.extend(lst)
            
    if other_list:
        doc.add_paragraph('BẢNG ĐÁP ÁN KHÁC').bold = True
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Câu'; hdr[1].text = 'Đáp án'; hdr[2].text = 'Câu'; hdr[3].text = 'Đáp án'
        total_qs = len(other_list)
        for i in range(0, total_qs, 2):
            row = table.add_row().cells
            row[0].text = str(other_list[i]['q']); row[1].text = str(other_list[i]['ans'])
            if i + 1 < total_qs:
                row[2].text = str(other_list[i+1]['q']); row[3].text = str(other_list[i+1]['ans'])
        doc.add_paragraph('')

def generate_mixed_exams(global_pools, structure, title_to_part, outro, merged_bytes, num_versions=4, q_sort_mode='random', shuffle_a=True, even_split=True, keep_color=False, manual_answers=None):
    zip_buffer = io.BytesIO()
    master_key_doc = Document()
    master_key_doc.add_heading('BẢNG ĐÁP ÁN TỔNG HỢP & TRUY VẾT NGUỒN', 0)

    with zipfile.ZipFile(zip_buffer, 'a', zipfile.ZIP_DEFLATED) as zip_file:
        for v in range(num_versions):
            ma_de = str(random.randint(100, 999))
            
            merged_stream = io.BytesIO(merged_bytes)
            new_doc = Document(merged_stream)
            body = new_doc._body._body
            
            sectPr = body.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr')
            if sectPr is not None: body.remove(sectPr)
            for child in list(body):
                if child.tag.endswith(('}p', '}tbl')): body.remove(child)
            
            ans_key = {}
            ans_by_part = {'part_1': [], 'part_2': [], 'part_3': [], 'other': []}
            source_info = []
            q_count = 1
            elements_to_keep = []
            
            for sec_template in structure:
                for p in sec_template['intro']: elements_to_keep.append(copy.deepcopy(p._element))
                title = sec_template['title']
                pool = global_pools.get(title, [])
                needed = sec_template['needed']
                
                actual_needed = min(needed, len(pool))
                randomize = (q_sort_mode == 'random')
                if even_split:
                    selected_items = allocate_questions_fairly(pool, actual_needed, randomize=randomize)
                else:
                    if randomize:
                        selected_items = random.sample(pool, actual_needed)
                    else:
                        selected_items = pool[:actual_needed]
                
                if q_sort_mode == 'random':
                    random.shuffle(selected_items)
                elif q_sort_mode == 'reverse':
                    selected_items.reverse()
                
                part_key = title_to_part.get(title)
                part_manual_ans = manual_answers.get(part_key, {}) if manual_answers else {}
                
                for item in selected_items:
                    src_name = item['src']
                    if item['type'] == 'q':
                        manual_ans_for_item = part_manual_ans.get(item.get('original_index'))
                        correct_label, els = process_question_block(item['data'], q_count, shuffle_a, keep_color, manual_ans=manual_ans_for_item)
                        elements_to_keep.extend(els)
                        ans_key[q_count] = correct_label
                        ans_info = {'q': q_count, 'ans': correct_label}
                        if part_key in ans_by_part: ans_by_part[part_key].append(ans_info)
                        else: ans_by_part['other'].append(ans_info)
                        
                        source_info.append(f'Câu {q_count}: Nguồn từ file [{src_name}]')
                        q_count += 1
                    elif item['type'] == 'g':
                        g_data = copy.deepcopy(item['data'])
                        for p in g_data['intro']: 
                            p.text = re.sub(r'^(Câu|Bài)\s+\d+', f'Câu {q_count}', p.text, flags=re.IGNORECASE)
                            elements_to_keep.append(copy.deepcopy(p._element))
                        
                        group_qs = g_data['questions']
                        sub_q_indices = item.get('original_indices', [])
                        
                        zipped_qs = list(zip(group_qs, sub_q_indices))
                        
                        if q_sort_mode == 'random':
                            random.shuffle(zipped_qs)
                        elif q_sort_mode == 'reverse':
                            zipped_qs.reverse()
                            
                        for q_blocks, orig_sub_idx in zipped_qs:
                            manual_ans_for_sub = part_manual_ans.get(orig_sub_idx)
                            correct_label, els = process_question_block(q_blocks, q_count, shuffle_a, keep_color, manual_ans=manual_ans_for_sub)
                            elements_to_keep.extend(els)
                            ans_key[q_count] = correct_label
                            ans_info = {'q': q_count, 'ans': correct_label}
                            if part_key in ans_by_part: ans_by_part[part_key].append(ans_info)
                            else: ans_by_part['other'].append(ans_info)
                            
                            source_info.append(f'Câu {q_count}: Nguồn từ file [{src_name}]')
                            q_count += 1
            
            for p in outro: elements_to_keep.append(copy.deepcopy(p._element))
            for el in elements_to_keep: body.append(el)
                
            try:
                new_doc.add_page_break()
                build_answer_tables(new_doc, ans_by_part, ma_de)
            except Exception:
                new_doc.add_paragraph('--- ĐÁP ÁN ---')
                for i in range(1, q_count): new_doc.add_paragraph(f'Câu {i}: {ans_key.get(i, "")}')
                    
            new_doc.add_paragraph('')
            new_doc.add_paragraph('📍 TRUY VẾT NGUỒN GỐC CÂU HỎI').bold = True
            for info in source_info: new_doc.add_paragraph(info)
                
            if sectPr is not None: body.append(sectPr)
            buf = io.BytesIO(); new_doc.save(buf)
            zip_file.writestr(f'De_{ma_de}.docx', buf.getvalue())
            
            master_key_doc.add_heading(f'MÃ ĐỀ: {ma_de}', level=1)
            try:
                build_answer_tables(master_key_doc, ans_by_part)
            except Exception: pass
            
            master_key_doc.add_paragraph('📍 Nguồn gốc câu hỏi:').bold = True
            for info in source_info: master_key_doc.add_paragraph(info, style='List Bullet')
            master_key_doc.add_page_break()

    zip_buffer.seek(0)
    master_buf = io.BytesIO(); master_key_doc.save(master_buf); master_buf.seek(0)
    return zip_buffer, master_buf


def normalize_clumped_choices(doc):
    import re
    from copy import deepcopy
    pattern = re.compile(r'(\s+)([A-Da-d]\*?[\.\)])(\s+)')
    
    for p in tuple(doc.paragraphs):
        text = p.text
        if not text.strip(): continue
        
        matches = list(pattern.finditer(text))
        if not matches: continue
        
        split_indices = [0] + [m.start(2) for m in matches] + [len(text)]
        
        for i in range(len(split_indices) - 1):
            start_idx = split_indices[i]
            end_idx = split_indices[i+1]
            
            if start_idx == end_idx: continue
            
            new_p = p.insert_paragraph_before()
            if p._p.pPr is not None: new_p._p.append(deepcopy(p._p.pPr))
            
            current_len = 0
            for run in p.runs:
                if not run.text:
                    if current_len >= start_idx and current_len < end_idx:
                        new_p._p.append(deepcopy(run._element))
                    continue
                    
                run_text = run.text
                run_len = len(run_text)
                
                overlap_start = max(current_len, start_idx)
                overlap_end = min(current_len + run_len, end_idx)
                
                if overlap_start < overlap_end:
                    rel_start = overlap_start - current_len
                    rel_end = overlap_end - current_len
                    segment_text = run_text[rel_start:rel_end]
                    
                    new_run = deepcopy(run._element)
                    for t in new_run.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                        t.text = segment_text
                    new_p._p.append(new_run)
                    
                current_len += run_len
                
        p._element.getparent().remove(p._element)

def parse_and_build_pools(files):
    merged_stream = merge_and_track_files(files)
    doc = Document(merged_stream)
    
    # Chuẩn hóa các phương án dính liền nhau (A. B. C. D. trên 1 dòng)
    normalize_clumped_choices(doc)
    
    sections, outro = extract_sections_xml(doc)
    
    global_pools = {} 
    structure = []
    
    title_to_part = {}
    part_idx = 0
    
    for sec in sections:
        title = sec['title']
        if title != 'CHUNG' and title not in title_to_part:
            match = re.search(r'PHẦN\s+([IVX]+)', title.upper())
            if match:
                roman = match.group(1)
                if roman == 'I': title_to_part[title] = 'part_1'
                elif roman == 'II': title_to_part[title] = 'part_2'
                elif roman == 'III': title_to_part[title] = 'part_3'
                else: 
                    part_idx += 1
                    title_to_part[title] = f'part_{part_idx}'
            else:
                part_idx += 1
                title_to_part[title] = f'part_{part_idx}'
            
        if title not in global_pools: global_pools[title] = []
        
        start_q_idx = sum(1 for item in global_pools[title] if item['type'] == 'q')
        start_q_idx += sum(len(item['data']['questions']) for item in global_pools[title] if item['type'] == 'g')
        
        current_q_idx = start_q_idx + 1
        
        for item in sec['items']:
            if item['type'] == 'q':
                item['original_index'] = current_q_idx
                current_q_idx += 1
            elif item['type'] == 'g':
                item['original_indices'] = list(range(current_q_idx, current_q_idx + len(item['data']['questions'])))
                current_q_idx += len(item['data']['questions'])
                
        global_pools[title].extend(sec['items'])
        if len(sec['items']) > 0 and sec['items'][0]['is_temp']:
            if not any(s['title'] == title for s in structure):
                structure.append({'title': title, 'intro': sec['intro'], 'needed': len(sec['items'])})
            else:
                for s in structure:
                    if s['title'] == title: s['needed'] += len(sec['items'])

    return global_pools, structure, title_to_part, outro, merged_stream.getvalue()

def generate_preview(global_pools, structure, title_to_part, manual_answers=None):
    preview_data = []
    
    q_count_mock = 1
    for sec_template in structure:
        title = sec_template['title']
        pool = global_pools.get(title, [])
        part_key = title_to_part.get(title)
        part_manual_ans = manual_answers.get(part_key, {}) if manual_answers else {}
        
        part_preview = {'title': title, 'questions': []}
        
        for item in pool:
            if item['type'] == 'q':
                manual_ans_for_item = part_manual_ans.get(item.get('original_index'))
                correct_label, els = process_question_block(item['data'], q_count_mock, shuffle_a=False, keep_color=True, manual_ans=manual_ans_for_item)
                
                full_text = '\n'.join(p.text for p in els if p.text.strip())
                part_preview['questions'].append({
                    'text': full_text, 
                    'answer': correct_label,
                    'orig_idx': item.get('original_index')
                })
                q_count_mock += 1
            elif item['type'] == 'g':
                g_data = copy.deepcopy(item['data'])
                intro_text = '\n'.join(p.text for p in g_data['intro'] if p.text.strip())
                
                group_qs = g_data['questions']
                sub_q_indices = item.get('original_indices', [])
                zipped_qs = list(zip(group_qs, sub_q_indices))
                
                for q_blocks, orig_sub_idx in zipped_qs:
                    manual_ans_for_sub = part_manual_ans.get(orig_sub_idx)
                    correct_label, els = process_question_block(q_blocks, q_count_mock, shuffle_a=False, keep_color=True, manual_ans=manual_ans_for_sub)
                    full_text = '\n'.join(p.text for p in els if p.text.strip())
                    
                    part_preview['questions'].append({
                        'text': f'[Nhóm dùng chung]\n{intro_text}\n{full_text}', 
                        'answer': correct_label,
                        'orig_idx': orig_sub_idx
                    })
                    q_count_mock += 1
                    
        preview_data.append(part_preview)
    return preview_data

def build_answer_tables(doc, ans_by_part, ma_de=None):
    import re
    if ma_de:
        doc.add_paragraph(f'BẢNG ĐÁP ÁN - MÃ ĐỀ {ma_de}').bold = True
        
    p1_list = ans_by_part.get('part_1', [])
    if p1_list:
        doc.add_paragraph('BẢNG ĐÁP ÁN PHẦN 1').bold = True
        chunk_size = 15
        for i in range(0, len(p1_list), chunk_size):
            chunk = p1_list[i:i+chunk_size]
            table = doc.add_table(rows=2, cols=len(chunk) + 1)
            table.style = 'Table Grid'
            table.cell(0,0).text = 'Câu'
            table.cell(1,0).text = 'Đáp án'
            for j, item in enumerate(chunk):
                table.cell(0, j+1).text = str(item['q'])
                table.cell(1, j+1).text = str(item['ans'])
        doc.add_paragraph('')

    p2_list = ans_by_part.get('part_2', [])
    if p2_list:
        doc.add_paragraph('BẢNG ĐÁP ÁN PHẦN 2').bold = True
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Câu'; hdr[1].text = 'a'; hdr[2].text = 'b'; hdr[3].text = 'c'; hdr[4].text = 'd'
        
        for item in p2_list:
            row = table.add_row().cells
            row[0].text = str(item['q'])
            ans_str = str(item['ans'])
            vals = {'a': '', 'b': '', 'c': '', 'd': ''}
            matches = re.findall(r'([a-d])\s*-\s*([ĐS])', ans_str, re.IGNORECASE)
            for k, v in matches: vals[k.lower()] = v.upper()
            row[1].text = vals['a']; row[2].text = vals['b']; row[3].text = vals['c']; row[4].text = vals['d']
        doc.add_paragraph('')

    p3_list = ans_by_part.get('part_3', [])
    if p3_list:
        doc.add_paragraph('BẢNG ĐÁP ÁN PHẦN 3').bold = True
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Câu'; hdr[1].text = 'Đáp án'
        for item in p3_list:
            row = table.add_row().cells
            row[0].text = str(item['q'])
            row[1].text = str(item['ans']).replace(';', '\n')
        doc.add_paragraph('')
        
    other_list = ans_by_part.get('other', [])
    for k, lst in ans_by_part.items():
        if k not in ['part_1', 'part_2', 'part_3', 'other']:
            other_list.extend(lst)
            
    if other_list:
        doc.add_paragraph('BẢNG ĐÁP ÁN KHÁC').bold = True
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Câu'; hdr[1].text = 'Đáp án'; hdr[2].text = 'Câu'; hdr[3].text = 'Đáp án'
        total_qs = len(other_list)
        for i in range(0, total_qs, 2):
            row = table.add_row().cells
            row[0].text = str(other_list[i]['q']); row[1].text = str(other_list[i]['ans'])
            if i + 1 < total_qs:
                row[2].text = str(other_list[i+1]['q']); row[3].text = str(other_list[i+1]['ans'])
        doc.add_paragraph('')

def build_answer_tables(doc, ans_by_part, ma_de=None):
    import re
    if ma_de:
        doc.add_paragraph(f'BẢNG ĐÁP ÁN - MÃ ĐỀ {ma_de}').bold = True
        
    p1_list = ans_by_part.get('part_1', [])
    if p1_list:
        doc.add_paragraph('BẢNG ĐÁP ÁN PHẦN 1').bold = True
        chunk_size = 15
        for i in range(0, len(p1_list), chunk_size):
            chunk = p1_list[i:i+chunk_size]
            table = doc.add_table(rows=2, cols=len(chunk) + 1)
            table.style = 'Table Grid'
            table.cell(0,0).text = 'Câu'
            table.cell(1,0).text = 'Đáp án'
            for j, item in enumerate(chunk):
                table.cell(0, j+1).text = str(item['q'])
                table.cell(1, j+1).text = str(item['ans'])
        doc.add_paragraph('')

    p2_list = ans_by_part.get('part_2', [])
    if p2_list:
        doc.add_paragraph('BẢNG ĐÁP ÁN PHẦN 2').bold = True
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Câu'; hdr[1].text = 'a'; hdr[2].text = 'b'; hdr[3].text = 'c'; hdr[4].text = 'd'
        
        for item in p2_list:
            row = table.add_row().cells
            row[0].text = str(item['q'])
            ans_str = str(item['ans'])
            vals = {'a': '', 'b': '', 'c': '', 'd': ''}
            matches = re.findall(r'([a-d])\s*-\s*([ĐS])', ans_str, re.IGNORECASE)
            for k, v in matches: vals[k.lower()] = v.upper()
            row[1].text = vals['a']; row[2].text = vals['b']; row[3].text = vals['c']; row[4].text = vals['d']
        doc.add_paragraph('')

    p3_list = ans_by_part.get('part_3', [])
    if p3_list:
        doc.add_paragraph('BẢNG ĐÁP ÁN PHẦN 3').bold = True
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Câu'; hdr[1].text = 'Đáp án'
        for item in p3_list:
            row = table.add_row().cells
            row[0].text = str(item['q'])
            row[1].text = str(item['ans']).replace(';', '\n')
        doc.add_paragraph('')
        
    other_list = ans_by_part.get('other', [])
    for k, lst in ans_by_part.items():
        if k not in ['part_1', 'part_2', 'part_3', 'other']:
            other_list.extend(lst)
            
    if other_list:
        doc.add_paragraph('BẢNG ĐÁP ÁN KHÁC').bold = True
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Câu'; hdr[1].text = 'Đáp án'; hdr[2].text = 'Câu'; hdr[3].text = 'Đáp án'
        total_qs = len(other_list)
        for i in range(0, total_qs, 2):
            row = table.add_row().cells
            row[0].text = str(other_list[i]['q']); row[1].text = str(other_list[i]['ans'])
            if i + 1 < total_qs:
                row[2].text = str(other_list[i+1]['q']); row[3].text = str(other_list[i+1]['ans'])
        doc.add_paragraph('')

def generate_mixed_exams(global_pools, structure, title_to_part, outro, merged_bytes, num_versions=4, q_sort_mode='random', shuffle_a=True, even_split=True, keep_color=False, manual_answers=None):
    zip_buffer = io.BytesIO()
    master_key_doc = Document()
    master_key_doc.add_heading('BẢNG ĐÁP ÁN TỔNG HỢP & TRUY VẾT NGUỒN', 0)

    with zipfile.ZipFile(zip_buffer, 'a', zipfile.ZIP_DEFLATED) as zip_file:
        for v in range(num_versions):
            ma_de = str(random.randint(100, 999))
            
            merged_stream = io.BytesIO(merged_bytes)
            new_doc = Document(merged_stream)
            body = new_doc._body._body
            
            sectPr = body.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr')
            if sectPr is not None: body.remove(sectPr)
            for child in list(body):
                if child.tag.endswith(('}p', '}tbl')): body.remove(child)
            
            ans_key = {}
            source_info = []
            q_count = 1
            elements_to_keep = []
            
            for sec_template in structure:
                for p in sec_template['intro']: elements_to_keep.append(copy.deepcopy(p._element))
                title = sec_template['title']
                pool = global_pools.get(title, [])
                needed = sec_template['needed']
                
                actual_needed = min(needed, len(pool))
                randomize = (q_sort_mode == 'random')
                if even_split:
                    selected_items = allocate_questions_fairly(pool, actual_needed, randomize=randomize)
                else:
                    if randomize:
                        selected_items = random.sample(pool, actual_needed)
                    else:
                        selected_items = pool[:actual_needed]
                
                if q_sort_mode == 'random':
                    random.shuffle(selected_items)
                elif q_sort_mode == 'reverse':
                    selected_items.reverse()
                
                part_key = title_to_part.get(title)
                part_manual_ans = manual_answers.get(part_key, {}) if manual_answers else {}
                
                for item in selected_items:
                    src_name = item['src']
                    if item['type'] == 'q':
                        manual_ans_for_item = part_manual_ans.get(item.get('original_index'))
                        correct_label, els = process_question_block(item['data'], q_count, shuffle_a, keep_color, manual_ans=manual_ans_for_item)
                        elements_to_keep.extend(els)
                        ans_key[q_count] = correct_label
                        source_info.append(f'Câu {q_count}: Nguồn từ file [{src_name}]')
                        q_count += 1
                    elif item['type'] == 'g':
                        g_data = copy.deepcopy(item['data'])
                        for p in g_data['intro']: 
                            p.text = re.sub(r'^(Câu|Bài)\s+\d+', f'Câu {q_count}', p.text, flags=re.IGNORECASE)
                            elements_to_keep.append(copy.deepcopy(p._element))
                        
                        group_qs = g_data['questions']
                        sub_q_indices = item.get('original_indices', [])
                        
                        zipped_qs = list(zip(group_qs, sub_q_indices))
                        
                        if q_sort_mode == 'random':
                            random.shuffle(zipped_qs)
                        elif q_sort_mode == 'reverse':
                            zipped_qs.reverse()
                            
                        for q_blocks, orig_sub_idx in zipped_qs:
                            manual_ans_for_sub = part_manual_ans.get(orig_sub_idx)
                            correct_label, els = process_question_block(q_blocks, q_count, shuffle_a, keep_color, manual_ans=manual_ans_for_sub)
                            elements_to_keep.extend(els)
                            ans_key[q_count] = correct_label
                            source_info.append(f'Câu {q_count}: Nguồn từ file [{src_name}]')
                            q_count += 1
            
            for p in outro: elements_to_keep.append(copy.deepcopy(p._element))
            for el in elements_to_keep: body.append(el)
                
            try:
                new_doc.add_page_break()
                new_doc.add_paragraph(f'BẢNG ĐÁP ÁN - MÃ ĐỀ {ma_de}').bold = True
                table = new_doc.add_table(rows=1, cols=4); table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Câu'; hdr_cells[1].text = 'Đáp án'; hdr_cells[2].text = 'Câu'; hdr_cells[3].text = 'Đáp án'
                total_qs = q_count - 1
                for i in range(1, total_qs + 1, 2):
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(i); row_cells[1].text = ans_key.get(i, '')
                    if i + 1 <= total_qs:
                        row_cells[2].text = str(i + 1); row_cells[3].text = ans_key.get(i + 1, '')
            except Exception:
                new_doc.add_paragraph('--- ĐÁP ÁN ---')
                for i in range(1, q_count): new_doc.add_paragraph(f'Câu {i}: {ans_key.get(i, "")}')
                    
            new_doc.add_paragraph('')
            new_doc.add_paragraph('📍 TRUY VẾT NGUỒN GỐC CÂU HỎI').bold = True
            for info in source_info: new_doc.add_paragraph(info)
                
            if sectPr is not None: body.append(sectPr)
            buf = io.BytesIO(); new_doc.save(buf)
            zip_file.writestr(f'De_{ma_de}.docx', buf.getvalue())
            
            master_key_doc.add_heading(f'MÃ ĐỀ: {ma_de}', level=1)
            try:
                table_mk = master_key_doc.add_table(rows=1, cols=4); table_mk.style = 'Table Grid'
                h_mk = table_mk.rows[0].cells
                h_mk[0].text = 'Câu'; h_mk[1].text = 'Đáp án'; h_mk[2].text = 'Câu'; h_mk[3].text = 'Đáp án'
                for i in range(1, q_count, 2):
                    r_mk = table_mk.add_row().cells
                    r_mk[0].text = str(i); r_mk[1].text = ans_key.get(i, '')
                    if i + 1 < q_count: r_mk[2].text = str(i + 1); r_mk[3].text = ans_key.get(i + 1, '')
            except Exception: pass
            
            master_key_doc.add_paragraph('📍 Nguồn gốc câu hỏi:').bold = True
            for info in source_info: master_key_doc.add_paragraph(info, style='List Bullet')
            master_key_doc.add_page_break()

    zip_buffer.seek(0)
    master_buf = io.BytesIO(); master_key_doc.save(master_buf); master_buf.seek(0)
    return zip_buffer, master_buf
