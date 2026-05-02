from docx import Document
from docx.shared import RGBColor
doc = Document()
p = doc.add_paragraph()
r = p.add_run('A.')
r.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

def is_correct_choice(paragraph):
    for run in paragraph.runs:
        if run.text.strip():
            rPr_xml = ''
            rPr_elems = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
            if rPr_elems:
                rPr_xml = rPr_elems[0].xml
                print('rPr_xml:', rPr_xml)
            if 'w:val="FF0000"' in rPr_xml:
                return True
    return False

print('Is correct:', is_correct_choice(p))
